from __future__ import annotations

import argparse
import logging
import os
import re
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import docx
from docx.document import Document
from docx.table import Table, _Cell
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn

import openpyxl
from openpyxl.workbook import Workbook
from openpyxl.worksheet.worksheet import Worksheet
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter

# Configure basic logging for the module
logger = logging.getLogger(__name__)

class DocxTableExtractor:
    """
    Extracts tables from a .docx file and exports them to a .xlsx workbook.
    """

    def __init__(self, docx_path: str | Path, quiet: bool = False):
        self.docx_path = Path(docx_path)
        if not self.docx_path.exists():
            raise FileNotFoundError(f"Input file not found: {self.docx_path}")

        self.doc = docx.Document(self.docx_path)
        self.workbook = openpyxl.Workbook()
        # Remove default sheet
        self.workbook.remove(self.workbook.active)
        self.quiet = quiet

    def process(self, output_path: str | Path | None = None, prefix: str = "Table") -> Path:
        """
        Process the document, extract tables, and save to Excel.
        """
        if not self.doc.tables:
            logger.warning(f"No tables found in {self.docx_path}.")

        for idx, table in enumerate(self.doc.tables):
            sheet_name = self._determine_sheet_name(table._tbl, idx, prefix)
            ws = self.workbook.create_sheet(title=sheet_name)
            self._process_table(table, ws)
            self._autofit_columns(ws)

        if not output_path:
            output_path = self.docx_path.with_suffix(".xlsx")
        else:
            output_path = Path(output_path)

        # If no tables were found, ensure at least one sheet exists to save valid xlsx
        if not self.workbook.sheetnames:
            self.workbook.create_sheet("Empty")

        self.workbook.save(output_path)
        logger.info(f"Successfully saved to {output_path}")
        return output_path

    def _sanitize_sheet_name(self, name: str) -> str:
        """
        Sanitize sheet name to conform to Excel rules:
        - Max 31 chars
        - No forbidden chars: \\ / ? * [ ] :
        """
        forbidden = r"[\\/\?\*\[\]:]"
        clean = re.sub(forbidden, "_", name).strip()
        if not clean:
            clean = "Table"
        return clean[:31]

    def _determine_sheet_name(self, tbl: CT_Tbl, idx: int, prefix: str) -> str:
        """
        Dynamically infer sheet name by walking up to 3 backward siblings.
        Fallback to prefix_idx if no valid text found or if name already exists.
        """
        name_candidate = None
        current = tbl
        for _ in range(3):
            prev = current.getprevious()
            if prev is None:
                break
            if prev.tag.endswith("p"):  # Paragraph w:p
                # Check for text in w:t
                texts = prev.xpath(".//w:t")
                text = "".join([t.text for t in texts if t.text]).strip()
                if text:
                    name_candidate = text
                    break
            current = prev

        if not name_candidate:
            name_candidate = f"{prefix}_{idx+1}"

        base_name = self._sanitize_sheet_name(name_candidate)
        final_name = base_name

        # Ensure absolute uniqueness
        counter = 1
        existing_names = self.workbook.sheetnames
        while final_name in existing_names:
            suffix = f" {counter}"
            # Adjust to fit max 31 chars with suffix
            max_base_len = 31 - len(suffix)
            final_name = f"{base_name[:max_base_len]}{suffix}"
            counter += 1

        return final_name

    def _get_cell_style(self, cell: _Cell) -> dict[str, Any]:
        """
        Extract style from the first run of the first paragraph in the cell.
        Returns a dict of style attributes.
        """
        style = {
            "bold": False,
            "italic": False,
            "underline": False,
            "size": None,
            "color": None,
            "alignment": "left"
        }

        # Try paragraph alignment
        if cell.paragraphs:
            para = cell.paragraphs[0]
            # Paragraph alignment
            align = para.alignment
            if align is not None:
                # Map docx alignment to openpyxl alignment
                if align == docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER:
                    style["alignment"] = "center"
                elif align == docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT:
                    style["alignment"] = "right"
                elif align == docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY:
                    style["alignment"] = "justify"

            # Find the first run with text to determine style
            for run in para.runs:
                if run.text.strip():
                    if run.bold:
                        style["bold"] = True
                    if run.italic:
                        style["italic"] = True
                    if run.underline:
                        style["underline"] = True
                    if run.font.size:
                        style["size"] = run.font.size.pt
                    if run.font.color and run.font.color.rgb:
                        # docx RGBColor is an object with hex property or str
                        style["color"] = str(run.font.color.rgb)
                    break

        return style

    def _process_table(self, table: Table, ws: Worksheet) -> None:
        """
        Process the table using grid attributes to correctly handle merges.
        """
        tbl = table._tbl
        # Parse the grid definition w:tblGrid
        grid = tbl.tblGrid
        if grid is None:
            grid_cols = len(table.columns)
        else:
            grid_cols = len(grid.gridCol_lst)

        # To handle vertical merges and track occupied matrix slots:
        # Matrix to hold (text, style, merge_width, merge_height)
        # However, it's easier to process XML rows, tracking vMerge state.

        # openpyxl uses 1-based indexing
        # Let's read row by row and map to a sparse grid.

        # Keep track of vertical spans: column_index -> (start_row, remaining_span, text, style)
        # Actually, in OpenXML w:vMerge, it has "restart" (starts merge) or None/"continue" (continues).
        # We need to map the visual cells to the w:tblGrid.

        # In python-docx, `table.rows` actually iterates over logical rows but can duplicate cells
        # that span horizontally or vertically.
        # But wait, python-docx duplicates cells for merged cells!
        # If A1 spans 3 cols, A1, B1, C1 in `table.row[0].cells` are the SAME cell object.
        # So our simplistic grid_matrix building over table.rows gets confused by duplicates.

        # Let's read from the XML directly.
        # w:tr -> w:tc

        grid_matrix: dict[tuple[int, int], dict[str, Any]] = {}

        for r_idx, tr in enumerate(tbl.tr_lst):
            c_idx = 0
            for tc in tr.tc_lst:
                # Find next available slot (might be filled by previous row's vMerge)
                while (r_idx, c_idx) in grid_matrix and grid_matrix[(r_idx, c_idx)]["is_merged_continuation"]:
                    c_idx += 1

                if c_idx >= grid_cols:
                    break

                # We need to construct a docx _Cell to reuse _get_cell_style and text extraction
                cell = _Cell(tc, table)

                tcPr = tc.tcPr
                if tcPr is not None:
                    gridSpan = tcPr.gridSpan
                    span_c = gridSpan.val if gridSpan is not None else 1

                    vMerge = tcPr.vMerge
                    vMerge_val = vMerge.val if vMerge is not None else None
                else:
                    span_c = 1
                    vMerge_val = None

                if vMerge_val is not None:
                    if vMerge_val == "restart":
                        text = cell.text.strip()
                        style = self._get_cell_style(cell)
                        grid_matrix[(r_idx, c_idx)] = {
                            "text": text,
                            "style": style,
                            "span_c": span_c,
                            "span_r": 1,
                            "is_merged_continuation": False
                        }
                        for i in range(1, span_c):
                            grid_matrix[(r_idx, c_idx + i)] = {"is_merged_continuation": True}
                    else: # continue
                        # Find restart cell above
                        restart_r_idx = r_idx - 1
                        while restart_r_idx >= 0:
                            if (restart_r_idx, c_idx) in grid_matrix and not grid_matrix[(restart_r_idx, c_idx)].get("is_merged_continuation", False):
                                grid_matrix[(restart_r_idx, c_idx)]["span_r"] += 1
                                break
                            elif (restart_r_idx, c_idx) in grid_matrix and grid_matrix[(restart_r_idx, c_idx)].get("is_merged_continuation", False):
                                # It might be a horizontal span continuation from the restart cell
                                found = False
                                for check_c_idx in range(c_idx, -1, -1):
                                    if (restart_r_idx, check_c_idx) in grid_matrix and not grid_matrix[(restart_r_idx, check_c_idx)].get("is_merged_continuation", False):
                                        grid_matrix[(restart_r_idx, check_c_idx)]["span_r"] += 1
                                        found = True
                                        break
                                if found:
                                    break
                            restart_r_idx -= 1

                        # Mark current slots as continuation (vMerge)
                        for i in range(span_c):
                            grid_matrix[(r_idx, c_idx + i)] = {"is_merged_continuation": True}
                else:
                    text = cell.text.strip()
                    style = self._get_cell_style(cell)
                    grid_matrix[(r_idx, c_idx)] = {
                        "text": text,
                        "style": style,
                        "span_c": span_c,
                        "span_r": 1,
                        "is_merged_continuation": False
                    }
                    for i in range(1, span_c):
                        grid_matrix[(r_idx, c_idx + i)] = {"is_merged_continuation": True}

                c_idx += span_c

        # Now write to openpyxl worksheet
        for r_idx in range(len(tbl.tr_lst)):
            for c_idx in range(grid_cols):
                if (r_idx, c_idx) in grid_matrix:
                    cell_data = grid_matrix[(r_idx, c_idx)]
                    if cell_data.get("is_merged_continuation"):
                        continue

                    xl_row = r_idx + 1
                    xl_col = c_idx + 1

                    text = cell_data["text"]
                    style = cell_data["style"]
                    span_r = cell_data["span_r"]
                    span_c = cell_data["span_c"]

                    xl_cell = ws.cell(row=xl_row, column=xl_col, value=text)

                    # Apply style
                    font_kwargs = {}
                    if style["bold"]: font_kwargs["bold"] = True
                    if style["italic"]: font_kwargs["italic"] = True
                    if style["underline"]: font_kwargs["underline"] = "single"
                    if style["size"]: font_kwargs["size"] = style["size"]
                    if style["color"]: font_kwargs["color"] = style["color"]

                    if font_kwargs:
                        xl_cell.font = Font(**font_kwargs)

                    align_kwargs = {
                        "wrap_text": True,
                        "vertical": "center",
                        "horizontal": style["alignment"]
                    }
                    xl_cell.alignment = Alignment(**align_kwargs)

                    # Apply merges
                    if span_r > 1 or span_c > 1:
                        ws.merge_cells(
                            start_row=xl_row, start_column=xl_col,
                            end_row=xl_row + span_r - 1, end_column=xl_col + span_c - 1
                        )

    def _autofit_columns(self, ws: Worksheet) -> None:
        """
        Dynamically adjust openpyxl column dimensions based on value lengths,
        with padding and a defensive maximum width cap.
        """
        MAX_WIDTH = 50.0
        PADDING = 2.0

        for col in ws.columns:
            max_len = 0
            # openpyxl MergedCell doesn't have column_letter, use the column index
            col_idx = col[0].column
            col_letter = get_column_letter(col_idx)

            for cell in col:
                if cell.value is not None:
                    # Ignore merged cells in length calculation as they can skew width
                    is_merged = any(
                        cell.coordinate in merged_range
                        for merged_range in ws.merged_cells.ranges
                    )
                    if not is_merged:
                        # Split by newline and get max line length
                        lines = str(cell.value).split("\n")
                        for line in lines:
                            if len(line) > max_len:
                                max_len = len(line)

            if max_len > 0:
                adjusted_width = max_len + PADDING
                # Cap at MAX_WIDTH
                ws.column_dimensions[col_letter].width = min(adjusted_width, MAX_WIDTH)

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Extract tables from a .docx file and export them to a .xlsx workbook."
    )
    parser.add_argument("input", help="Path to the input .docx file")
    parser.add_argument("-o", "--output", help="Path to the output .xlsx file (optional)", default=None)
    parser.add_argument("-p", "--prefix", help="Prefix for sheet names if no heading is found", default="Table")
    parser.add_argument("-q", "--quiet", help="Suppress output (except errors)", action="store_true")

    args = parser.parse_args()

    # Configure logging based on quiet flag
    log_level = logging.WARNING if args.quiet else logging.INFO
    logging.basicConfig(
        level=log_level,
        format="%(levelname)s: %(message)s"
    )

    try:
        extractor = DocxTableExtractor(args.input, quiet=args.quiet)
        extractor.process(output_path=args.output, prefix=args.prefix)
        sys.exit(0)
    except Exception as e:
        logger.error(f"Failed to process document: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
