import docx

doc = docx.Document()

# Paragraph to act as sheet name
doc.add_heading('My Test Data')

# Create a 3x3 table
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'

# Merge top row (horizontal merge)
cell1 = table.cell(0, 0)
cell1.text = "Header 1"
cell2 = table.cell(0, 1)
cell3 = table.cell(0, 2)
# We don't have direct merge in docx API easily? Wait, we do:
cell1.merge(cell2).merge(cell3)
cell1.text = "Merged Top Row"

# Add some text
table.cell(1, 0).text = "Row 2 Col 1"
table.cell(1, 1).text = "Row 2 Col 2"
table.cell(1, 2).text = "Row 2 Col 3"

# Merge first col of bottom rows (vertical merge)
cell_v1 = table.cell(1, 0)
cell_v2 = table.cell(2, 0)
cell_v1.merge(cell_v2)
cell_v1.text = "Merged V Col"

table.cell(2, 1).text = "Row 3 Col 2"
table.cell(2, 2).text = "Row 3 Col 3"

doc.save("test_table.docx")
print("Saved test_table.docx")
