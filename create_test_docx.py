from docx import Document

doc = Document()
doc.add_heading('Test Document', 0)

table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = 'A1'
table.cell(0, 1).text = 'B1'
table.cell(1, 0).text = 'A2'
table.cell(1, 1).text = 'B2'

table2 = doc.add_table(rows=3, cols=3)
for r in range(3):
    for c in range(3):
        table2.cell(r, c).text = f'R{r}C{c}'

doc.save('test.docx')
